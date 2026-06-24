from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path

DOC11 = Path('Doc11_filled.docx')
DOC12 = Path('Doc12_filled.docx')

def font(run, size=11, bold=False):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold

def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run(text)
    font(r, 18, True)
    p.paragraph_format.space_after = Pt(12)

def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 13, True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 11)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    font(r, 11)
    p.paragraph_format.space_after = Pt(3)

def table(doc, rows):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = 'Table Grid'
    for i, text in enumerate(rows[0]):
        r = t.rows[0].cells[i].paragraphs[0].add_run(text)
        font(r, 10.5, True)
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            r = cells[i].paragraphs[0].add_run(text)
            font(r, 10.5)
    doc.add_paragraph()

# Word1
_doc = Document()
title(_doc, '第五人格电竞社团宣传网站项目说明')
heading(_doc, '一、项目简介')
para(_doc, '本项目是一个 Vue2 制作的社团宣传网站，主题为“江西财经职业学院第五人格电竞社团”。网站主要用于展示社团介绍、社团活动、赛事回放和招新报名。')
heading(_doc, '二、页面功能')
bullet(_doc, '社团首页：展示社团名称、宣传文字、社团照片、社团特色和加入入口。')
bullet(_doc, '社团活动：展示社团常见活动，并使用嵌套路由切换活动概览和活动相册。')
bullet(_doc, '赛事回放：展示第五人格赛事回放入口，点击后可跳转到相关视频搜索页面。')
bullet(_doc, '加入我们：提供报名表单和招新二维码，方便同学提交报名意向。')
heading(_doc, '三、实现方式')
bullet(_doc, '使用 Vue2 脚手架创建项目。')
bullet(_doc, '使用 vue-router 实现页面切换，并在活动页面中使用嵌套路由。')
bullet(_doc, '将页面拆分为 views、components、api 等文件夹，避免代码全部写在 App.vue 中。')
bullet(_doc, '使用 axios 读取 public/api/club.json，模拟后端接口数据。')
bullet(_doc, '报名表单使用 v-model 获取输入内容，并用 localStorage 模拟保存。')
heading(_doc, '四、使用到的关键技术')
table(_doc, [
    ['技术', '作用'],
    ['Vue2', '完成页面组件化开发'],
    ['vue-router', '实现前端路由和嵌套路由'],
    ['axios', '访问模拟后端数据'],
    ['v-for / v-bind / v-on / v-model', '实现列表渲染、属性绑定、事件绑定和表单双向绑定'],
    ['CSS Grid / Flex', '完成页面布局和响应式适配']
])
heading(_doc, '五、项目特点')
bullet(_doc, '主题明确，围绕第五人格电竞社团展开。')
bullet(_doc, '页面较完整，包含展示、活动、视频入口和报名功能。')
bullet(_doc, '使用红黑金配色，整体风格偏电竞宣传。')
bullet(_doc, '后期可以继续接入真实后端接口。')
_doc.save(DOC11)

# Word2
_doc = Document()
title(_doc, '小组成员主要工作说明')
heading(_doc, '一、小组情况')
para(_doc, '本组共有 8 名成员，其中我担任组长。项目制作过程中，小组成员分别负责页面设计、功能实现、素材整理、测试和文档编写等工作。')
heading(_doc, '二、成员分工')
table(_doc, [
    ['成员', '主要工作'],
    ['组长（本人）', '确定项目主题，搭建项目结构，配置路由，整合页面，完成最终调试和文档整理。'],
    ['成员1', '负责社团首页内容和导航栏部分。'],
    ['成员2', '负责社团活动页面内容。'],
    ['成员3', '负责活动相册和图片素材整理。'],
    ['成员4', '负责赛事回放页面和视频链接整理。'],
    ['成员5', '负责加入我们页面和报名表单。'],
    ['成员6', '负责接口数据文件和 axios 数据读取。'],
    ['成员7', '负责页面样式优化、测试和截图整理。']
])
heading(_doc, '三、协作过程')
bullet(_doc, '先确定网站主题和页面结构。')
bullet(_doc, '再按页面和功能进行分工。')
bullet(_doc, '开发完成后统一整合代码和素材。')
bullet(_doc, '最后进行运行测试、页面调整和文档整理。')
heading(_doc, '四、总结')
para(_doc, '通过本次项目，小组成员对 Vue2 项目开发、组件化、前端路由、接口访问和页面样式设计有了更深入的理解，也提高了团队协作和项目整合能力。')
_doc.save(DOC12)
print('ok')
