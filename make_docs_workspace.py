from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

DOC11 = Path('Doc11_filled.docx')
DOC12 = Path('Doc12_filled.docx')

SCARLET = 'A61C28'
NIGHT = '0A0A0E'
GOLD = 'C8A460'
GRAY = '666677'
LIGHT = 'F4F4F6'
BORDER = 'DADCE0'


def set_font(run, size=10.5, bold=False, color='222222'):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color='222222', size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width * 1440)))
            tc_w.set(qn('w:type'), 'dxa')


def base_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    normal = doc.styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(10.5)
    return doc


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=21, bold=True, color=NIGHT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=10.5, color=GRAY)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(12)
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SCARLET)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=15 if level == 1 else 12.5, bold=True, color=SCARLET if level == 1 else NIGHT)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=10.5, color='222222')
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.5, color='222222')
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], SCARLET)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color='FFFFFF', size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, color='222222', size=9.5 if len(str(val)) > 55 else 10)
            if i == 0:
                set_cell_shading(cells[i], LIGHT)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def create_doc11():
    doc = base_doc()
    add_title(doc, '江西财经职业学院第五人格电竞社团宣传网站项目说明', 'Vue2 前端课程期末项目说明文档')

    add_heading(doc, '一、项目概述')
    add_para(doc, '本项目以“江西财经职业学院第五人格电竞社团”为主题，制作一个用于社团宣传、活动展示、赛事回放推荐和招新报名的 Vue2 单页应用。网站整体采用猩红、暗夜黑、金属灰、暗金等电竞风格配色，突出第五人格社团的竞技氛围和校园社团属性。')

    add_heading(doc, '二、页面功能说明')
    add_table(doc, ['页面', '主要功能'], [
        ('社团首页', '展示社团名称、宣传标语、社团照片、招新入口、赛事回放入口、社团特点和成员方向。'),
        ('社团活动', '使用嵌套路由展示活动概览和活动相册，包含新生体验局、校内水友赛、赛事观赛夜等内容。'),
        ('赛事回放', '展示第五人格职业赛、IVL、COA 等回放入口，并通过外部链接跳转到 B 站检索结果。'),
        ('加入我们', '提供报名意向表单，填写姓名、专业班级、加入方向和留言，并展示招新二维码。')
    ], [1.35, 5.15])

    add_heading(doc, '三、实现方式')
    for item in [
        '项目基于 Vue2 脚手架创建，入口文件为 main.js，主组件为 App.vue。',
        '使用 vue-router@3 实现前端路由，包含首页、社团活动、赛事回放、加入我们等页面。',
        '社团活动页使用父子嵌套路由，父页面 ActivitiesView 内部通过 router-view 展示活动概览和活动相册。',
        '页面内容拆分为 views、components、api、data 等目录，避免所有代码集中在 App.vue 中。',
        '使用 axios 访问 public/api/club.json，模拟后端接口返回社团数据、活动数据、赛事链接和图片路径。',
        '报名表单通过 v-model 收集用户输入，通过组件事件向父组件提交，并使用 localStorage 模拟后端保存。'
    ]:
        add_bullet(doc, item)

    add_heading(doc, '四、关键技术')
    add_table(doc, ['技术点', '项目中的应用'], [
        ('Vue2 组件化', '将导航、标题、图片展示、统计卡片、赛事卡片、报名表单拆分为独立组件。'),
        ('Vue Router', '使用 hash 模式实现单页应用页面切换，并完成活动页嵌套路由。'),
        ('Axios 接口访问', '通过 axios 请求本地 JSON 文件，模拟后端接口数据读取。'),
        ('Vue 指令', '使用 v-for 渲染列表，v-bind 绑定属性，v-on 绑定事件，v-model 实现表单双向绑定。'),
        ('组件通信', '子组件通过 $emit 向父组件传递导航和表单提交事件。'),
        ('CSS 响应式布局', '使用 Grid、Flex、媒体查询、clip-path、渐变背景等实现电竞风格界面。')
    ], [1.55, 4.95])

    add_heading(doc, '五、项目特色')
    for item in [
        '主题鲜明：围绕第五人格电竞社团展开，内容包含社团、活动、赛事和招新。',
        '结构完整：具备多个页面、嵌套路由、组件复用、接口访问和表单交互。',
        '视觉统一：采用猩红、暗夜黑、金属灰、暗金作为主视觉配色，增强电竞宣传感。',
        '便于扩展：后续可将本地 JSON 接口替换为真实后端接口，也可继续补充活动照片和报名数据库。'
    ]:
        add_bullet(doc, item)

    add_heading(doc, '六、运行与提交说明')
    for item in [
        '开发运行命令：npm run serve。',
        '项目构建命令：npm run build。',
        '提交项目文件夹时应删除 node_modules 文件夹，保留 src、public、package.json 等核心文件。',
        '项目使用的图片素材放置在 public/images 目录中，接口模拟数据放置在 public/api/club.json。'
    ]:
        add_bullet(doc, item)
    doc.save(DOC11)


def create_doc12():
    doc = base_doc()
    add_title(doc, '江西财经职业学院第五人格电竞社团宣传网站小组分工', '期末项目小组成员主要工作说明')

    add_heading(doc, '一、小组基本情况')
    add_para(doc, '本项目由 8 名同学共同完成，其中 1 名组长负责整体规划、进度协调和最终整合，7 名成员分别参与页面制作、组件开发、资料整理、样式设计、接口数据、测试优化和文档整理等工作。')

    add_heading(doc, '二、成员主要分工')
    add_table(doc, ['成员', '主要负责内容', '具体工作'], [
        ('组长（本人）', '整体规划与最终整合', '确定项目主题和页面结构，统筹 Vue2 项目搭建、路由配置、页面整合、最终调试和提交材料整理。'),
        ('成员1', '首页与导航模块', '负责社团首页内容整理、顶部导航栏、社团名称展示、首页宣传文案和入口按钮设计。'),
        ('成员2', '社团活动页面', '负责活动概览页内容，包括新生体验局、校内水友赛、赛事观赛夜等活动模块。'),
        ('成员3', '活动相册与图片素材', '负责收集和整理社团合照、游戏截图、活动海报、招新二维码等图片，并配合页面进行素材替换。'),
        ('成员4', '赛事回放页面', '负责整理第五人格 IVL、COA 等赛事回放入口，并制作赛事回放卡片内容。'),
        ('成员5', '加入我们表单', '负责报名表单字段设计、v-model 双向绑定、表单提交预览和招新二维码展示。'),
        ('成员6', '接口与数据维护', '负责 public/api/club.json 模拟后端数据维护，以及 axios 请求方法和本地存储逻辑。'),
        ('成员7', '样式优化与测试文档', '负责页面配色、响应式适配、构建测试、截图整理和项目说明文档初稿。')
    ], [1.15, 1.55, 3.8])

    add_heading(doc, '三、协作流程')
    for item in [
        '前期讨论项目主题，确定以“江西财经职业学院第五人格电竞社团宣传网站”为方向。',
        '组长划分页面和模块任务，各成员分别完成对应页面内容、素材整理和功能实现。',
        '开发过程中统一项目结构，按照 views、components、api、data 等目录组织代码。',
        '后期集中进行页面联调，检查路由跳转、表单提交、图片显示、响应式布局和构建结果。',
        '最终由组长整理项目文件、Word 文档、页面截图和提交说明。'
    ]:
        add_bullet(doc, item)

    add_heading(doc, '四、组长工作总结')
    add_para(doc, '作为组长，我主要负责项目整体方向把控、技术选型、项目结构设计、核心页面整合、路由和接口联调，以及最终提交材料的整理。通过本次项目，小组成员进一步熟悉了 Vue2 组件化开发、前端路由、接口数据访问和页面样式设计，也锻炼了团队协作与项目整合能力。')

    add_heading(doc, '五、后续可完善方向')
    for item in [
        '将模拟 JSON 接口替换为真实后端接口，实现报名信息数据库保存。',
        '继续补充更多社团活动照片、校内赛视频和成员介绍。',
        '增加后台管理页面，用于维护赛事回放、社团活动和报名名单。',
        '根据移动端浏览效果继续优化手机端布局。'
    ]:
        add_bullet(doc, item)
    doc.save(DOC12)


create_doc11()
create_doc12()
print(str(DOC11))
print(str(DOC12))

